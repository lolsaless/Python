import os
from tkinter import Tk
from tkinter.filedialog import askopenfilename
from docx import Document
from openpyxl import load_workbook


def find_and_replace_text(doc, find_text, replace_text):
    for paragraph in doc.paragraphs:
        if find_text in paragraph.text:
            paragraph.text = paragraph.text.replace(find_text, replace_text)


def find_and_replace_email(doc, find_email, replace_email):
    for paragraph in doc.paragraphs:
        if find_email in paragraph.text:
            paragraph.text = paragraph.text.replace(find_email, replace_email)


def select_files():
    Tk().withdraw()
    excel_file = askopenfilename(title="Select Excel File", filetypes=[("Excel Files", "*.xlsx")])
    word_file = askopenfilename(title="Select Word File", filetypes=[("Word Files", "*.docx")])
    return excel_file, word_file


def process_files():
    excel_file, word_file = select_files()

    if not excel_file or not word_file:
        print("File selection cancelled.")
        return

    wb = load_workbook(excel_file)
    sheet = wb.active

    for row in sheet.iter_rows(min_row=2, values_only=True):
        name, age, gender, email = row

        doc = Document(word_file)
        find_and_replace_text(doc, "홍길동", name)
        find_and_replace_text(doc, "20", str(age))
        find_and_replace_text(doc, "여", gender)
        find_and_replace_email(doc, "홍길동@kakao.com", email)

        output_folder = "수료증"
        os.makedirs(output_folder, exist_ok=True)
        output_filename = os.path.join(output_folder, f"{name}.docx")
        doc.save(output_filename)

    print("Processing complete.")


process_files()
