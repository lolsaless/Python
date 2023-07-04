import os
import tkinter as tk
from tkinter import filedialog
from docx import Document
from openpyxl import load_workbook


def find_and_replace_text(doc, find_text, replace_text):
    for paragraph in doc.paragraphs:
        if find_text in paragraph.text:
            paragraph.text = paragraph.text.replace(find_text, replace_text)

def find_and_replace_email(doc, find_email, replace_email):
    for paragraph in doc.paragraphs:
        if find_email in paragraph.text:
            paragraph.text = paragraph.text.replace(find_email, replace_email)

def select_excel_file():
    root = tk.Tk()
    root.withdraw()
    excel_path = filedialog.askopenfilename(filetypes=[("Excel Files", "*.xlsx")])
    return excel_path

def select_word_file():
    root = tk.Tk()
    root.withdraw()
    word_path = filedialog.askopenfilename(filetypes=[("Word Files", "*.docx")])
    return word_path

def main():
    excel_file = select_excel_file()
    word_file = select_word_file()

    if excel_file and word_file:
        wb = load_workbook(excel_file)
        sheet = wb.active

        for row in sheet.iter_rows(min_row=2, values_only=True):
            name = row[0]
            age = str(row[1])
            gender = row[2]
            email = row[3]

            doc = Document(word_file)
            find_and_replace_text(doc, "홍길동", name)
            find_and_replace_text(doc, "20", age)
            find_and_replace_text(doc, "여", gender)
            find_and_replace_email(doc, "홍길동@kakao.com", email)

            output_folder = "수료증"
            os.makedirs(output_folder, exist_ok=True)
            output_file = os.path.join(output_folder, f"{name}.docx")
            doc.save(output_file)

        print("작업이 완료되었습니다.")
    else:
        print("파일 선택이 취소되었습니다.")


if __name__ == "__main__":
    main()
