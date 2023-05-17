import os
import tkinter as tk
from tkinter import filedialog
from docx import Document
import openpyxl

def replace_text(doc, old_text, new_text):
    for paragraph in doc.paragraphs:
        if old_text in paragraph.text:
            paragraph.text = paragraph.text.replace(old_text, new_text)
    
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if old_text in cell.text:
                    cell.text = cell.text.replace(old_text, new_text)

def replace_email(doc, old_email, new_email):
    for paragraph in doc.paragraphs:
        if old_email in paragraph.text:
            paragraph.text = paragraph.text.replace(old_email, new_email)
    
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if old_email in cell.text:
                    cell.text = cell.text.replace(old_email, new_email)

def select_file():
    excel_filepath = filedialog.askopenfilename(filetypes=[("Excel files", "*.xlsx")])
    word_filepath = filedialog.askopenfilename(filetypes=[("Word files", "*.docx")])
    if excel_filepath and word_filepath:
        wb = openpyxl.load_workbook(excel_filepath)
        sheet = wb.active
        
        for row in sheet.iter_rows(min_row=2, values_only=True):
            name = row[0]
            age = row[1]
            gender = row[2]
            email = row[3]
            
            doc = Document(word_filepath)
            
            replace_text(doc, "홍길동", name)
            replace_text(doc, "20", str(age))
            replace_text(doc, "여", gender)
            replace_email(doc, "홍길동@kakao.com", email)
            
            output_folder = "수료증"
            if not os.path.exists(output_folder):
                os.makedirs(output_folder)
            
            new_filepath = os.path.join(output_folder, f"{name}.docx")
            doc.save(new_filepath)
            print("변경된 파일이 저장되었습니다:", new_filepath)

root = tk.Tk()
root.title("Excel 파일 선택")
root.geometry("200x100")

button = tk.Button(root, text="파일 선택", command=select_file)
button.pack(pady=20)

root.mainloop()
